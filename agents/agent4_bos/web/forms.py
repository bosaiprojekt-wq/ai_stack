
from fastapi import FastAPI, Form, Request
from fastapi.responses import HTMLResponse, JSONResponse
import os
from datetime import datetime
from docx import Document
import time

# Import from core modules
from core.config import SPECIAL_CASES_PATH
from core.document_ingestor import document_ingestor

# Create sub-app
form_app = FastAPI(title="Form Application", description="HTML form for adding cases")

def save_form_as_docx(title: str, author: str, description: str, solution: str, notes: str = "") -> dict:
    """Save form response as DOCX file - watcher zajmie się Qdrantem"""
    try:
        # Create directory if it doesn't exist
        os.makedirs(SPECIAL_CASES_PATH, exist_ok=True)
        
        # Generate filename with timestamp
        date_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in title)[:50]
        filename = f"form_{date_str}_{safe_title}.docx"
        filepath = os.path.join(SPECIAL_CASES_PATH, filename)
        
        # Create DOCX document
        doc = Document()
        
        # Add title
        doc.add_heading(f'Przypadek: {title}', 0)
        
        # Add metadata
        doc.add_paragraph(f'Autor: {author}')
        doc.add_paragraph(f'Data utworzenia: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'ID: FORM-{int(time.time())}')
        doc.add_paragraph('-' * 50)
        
        # Add description
        doc.add_heading('Opis przypadku:', level=1)
        doc.add_paragraph(description)
        
        # Add solution
        doc.add_heading('Rozwiązanie/procedura:', level=1)
        doc.add_paragraph(solution)
        
        # Add notes if provided
        if notes.strip():
            doc.add_heading('Uwagi dodatkowe:', level=1)
            doc.add_paragraph(notes)
        
        # Add footer
        doc.add_paragraph('-' * 50)
        doc.add_paragraph('Wygenerowano przez Agent4 BOS System')
        
        # Save document
        doc.save(filepath)
        print(f"✓ DOCX saved: {filename}")
        
        return {
            "filename": filename,
            "filepath": filepath,
            "docx_saved": True,
            "docx_size": os.path.getsize(filepath)
        }
        
    except Exception as e:
        print(f"✗ Error in save_form_as_docx: {e}")
        import traceback
        traceback.print_exc()
        return {
            "filename": None,
            "error": str(e),
            "docx_saved": False
        }


# FORM_HTML - bez zmian
FORM_HTML = """
<!DOCTYPE html>
<html>
<head>
    <title>Agent4 BOS - Formularz przypadku</title>
    <style>
        body { 
            font-family: Arial, sans-serif; 
            max-width: 800px; 
            margin: 40px auto; 
            padding: 20px; 
        }
        .container { 
            border: 1px solid #ccc; 
            padding: 30px; 
            border-radius: 8px; 
            background: #f9f9f9; 
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        h1 { 
            color: #2c3e50; 
            text-align: center; 
            margin-bottom: 30px;
        }
        label { 
            display: block; 
            margin-top: 15px; 
            font-weight: bold; 
            color: #34495e;
        }
        input[type="text"], textarea { 
            width: 100%; 
            padding: 10px; 
            margin-top: 5px; 
            border: 1px solid #ddd; 
            border-radius: 4px;
            font-size: 14px;
        }
        textarea { 
            height: 120px; 
            resize: vertical; 
        }
        .submit-btn { 
            margin-top: 25px; 
            padding: 12px 30px; 
            background: #3498db; 
            color: white; 
            border: none; 
            border-radius: 4px; 
            cursor: pointer; 
            font-size: 16px;
            display: block;
            width: 100%;
        }
        .submit-btn:hover { 
            background: #2980b9; 
        }
        .message {
            padding: 10px;
            margin: 10px 0;
            border-radius: 4px;
            display: none;
        }
        .success { background: #d4edda; color: #155724; border: 1px solid #c3e6cb; }
        .error { background: #f8d7da; color: #721c24; border: 1px solid #f5c6cb; }
        .info {
            background: #e8f4fd;
            padding: 15px;
            border-radius: 4px;
            margin-bottom: 20px;
            border-left: 4px solid #3498db;
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>📝 Formularz przypadku</h1>
        
        <div class="info">
            <strong>Informacja:</strong> Wypełnij formularz aby dodać nowy przypadek do bazy wiedzy.
            Po zapisaniu, plik DOCX zostanie automatycznie zaimportowany przez system.
        </div>
        
        <div id="message" class="message"></div>
        
        <form id="caseForm">
            <label>Tytuł przypadku:</label>
            <input type="text" name="Tytul" placeholder="Np. Urlop dziekański - procedura" required>
            
            <label>Autor (osoba dodająca):</label>
            <input type="text" name="Autor" placeholder="Twoje imię/nazwa" required>
            
            <label>Opis przypadku:</label>
            <textarea name="Opis" placeholder="Szczegółowy opis sytuacji, problemu..." required></textarea>
            
            <label>Rozwiązanie/procedura:</label>
            <textarea name="Rozwiazanie" placeholder="Proponowane rozwiązanie, kroki postępowania..." required></textarea>
            
            <label>Uwagi dodatkowe (opcjonalne):</label>
            <textarea name="Uwagi" placeholder="Dodatkowe informacje, komentarze..."></textarea>
            
            <button type="submit" class="submit-btn">💾 Zapisz przypadek (DOCX + auto-import)</button>
        </form>
        
        <div style="margin-top: 20px; text-align: center;">
            <a href="/" style="color: #3498db; text-decoration: none;">← Powrót do głównej strony</a>
        </div>
    </div>

    <script>
        document.getElementById('caseForm').addEventListener('submit', async function(e) {
            e.preventDefault();
            
            const formData = new FormData(this);
            const submitBtn = this.querySelector('button[type="submit"]');
            const originalText = submitBtn.textContent;
            
            submitBtn.textContent = 'Zapisywanie...';
            submitBtn.disabled = true;
            
            try {
                const response = await fetch('/form/submit', {
                    method: 'POST',
                    body: formData
                });
                
                const result = await response.json();
                
                if (result.status === 'success') {
                    let message = '✅ Przypadek został pomyślnie zapisany jako DOCX!';
                    if (result.docx_file) {
                        message += `\\n📄 Plik: ${result.docx_file}`;
                    }
                    message += '\\n\\n⏳ System automatycznie zaimportuje go do Qdrant za chwilę.';
                    showMessage('success', message);
                    this.reset();
                } else {
                    showMessage('error', '❌ Błąd: ' + (result.message || 'Nie udało się zapisać przypadku'));
                }
            } catch (error) {
                showMessage('error', '❌ Błąd połączenia: ' + error.message);
            } finally {
                submitBtn.textContent = originalText;
                submitBtn.disabled = false;
            }
        });
        
        function showMessage(type, text) {
            const msgDiv = document.getElementById('message');
            msgDiv.className = 'message ' + type;
            msgDiv.textContent = text;
            msgDiv.style.display = 'block';
            
            setTimeout(() => {
                msgDiv.style.display = 'none';
            }, 8000);
        }
    </script>
</body>
</html>
"""


@form_app.get("/")
async def get_form(request: Request):
    """Serve the HTML form"""
    return HTMLResponse(FORM_HTML)


@form_app.post("/submit")
async def submit_case(
    Tytul: str = Form(...),
    Autor: str = Form(...),
    Opis: str = Form(...),
    Rozwiazanie: str = Form(...),
    Uwagi: str = Form("")
):
    """Handle form submission - saves as DOCX, watcher handles Qdrant"""
    try:
        print(f"\n{'='*60}")
        print("FORM SUBMISSION RECEIVED")
        print(f"Title: {Tytul}")
        print(f"Author: {Autor}")
        print('='*60)
        
        result = save_form_as_docx(Tytul, Autor, Opis, Rozwiazanie, Uwagi)
        
        if not result.get("docx_saved"):
            return JSONResponse({
                "status": "error",
                "message": f"Błąd zapisu DOCX: {result.get('error', 'Unknown error')}"
            }, status_code=500)

        
        print(f"\n FORM SUBMISSION COMPLETE")
        print(f"   DOCX saved: {result['filename']}")
        print(f"   Waiting for auto-ingestion...")
        print('='*60)
        
        return JSONResponse({
            "status": "success",
            "message": "Przypadek zapisany jako DOCX. System zaimportuje go automatycznie.",
            "docx_file": result["filename"],
            "details": {
                "docx_size": result.get("docx_size", 0)
            }
        })
        
    except Exception as e:
        print(f"\nFORM SUBMISSION ERROR: {e}")
        import traceback
        traceback.print_exc()
        
        return JSONResponse({
            "status": "error",
            "message": f"Błąd systemu: {str(e)}"
        }, status_code=500)