#imports
import os
import time
import hashlib
import uuid
from pathlib import Path
from typing import List, Dict, Any
import docx
from PyPDF2 import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter

#class: DocumentProcessor - handles document processing (text extraction, chunking, metadata)
class DocumentProcessor:
    def __init__(self, chunk_size=1000, chunk_overlap=200):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", "? ", "! ", " ", ""]
        )

    #method: extract text from different file formats
    def extract_text(self, filepath: str) -> str:
        """Extract text from different file formats"""
        filepath = Path(filepath)
        
        if filepath.suffix.lower() == '.docx':
            return self._extract_from_docx(filepath)
        elif filepath.suffix.lower() == '.pdf':
            return self._extract_from_pdf(filepath)
        elif filepath.suffix.lower() == '.txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file format: {filepath.suffix}")
        
    #method: extract text from DOCX file
    def _extract_from_docx(self, filepath: Path) -> str:
        """Extract text from DOCX file"""
        doc = docx.Document(filepath)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n\n".join(text_parts)
    
    #method: extract text from PDF file
    def _extract_from_pdf(self, filepath: Path) -> str:
        """Extract text from PDF file"""
        reader = PdfReader(filepath)
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)
        
        return "\n\n".join(text_parts)
    
    #method: process file into chunks with metadata
    def process_file(self, filepath: str) -> List[Dict[str, Any]]:
        """Process a file into chunks with metadata"""
        filepath = Path(filepath)
        
        # Extract text
        try:
            text = self.extract_text(filepath)
        except Exception as e:
            print(f"Error extracting text from {filepath}: {e}")
            return []
        
        if not text.strip():
            print(f"No text content in {filepath}")
            return []
        
        # Split into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create records
        records = []
        
        for i, chunk in enumerate(chunks):
            # Create unique ID
            chunk_id = str(uuid.uuid4())
            
            # Determine category from folder structure
            category = self._determine_category(str(filepath))
            
            record = {
                "id": chunk_id,
                "text": chunk,
                "metadata": {
                    "source_file": str(filepath),
                    "filename": filepath.name,
                    "file_extension": filepath.suffix.lower(),
                    "file_size": filepath.stat().st_size,
                    "file_hash": self._calculate_file_hash(filepath),
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "category": category,
                    "ingestion_time": time.time(),
                    "last_modified": filepath.stat().st_mtime
                }
            }
            records.append(record)
        
        return records
    
    #method: determine category from file path
    def _determine_category(self, filepath: str) -> str:
        """Determine category from file path"""
        filepath_lower = filepath.lower()
        
        if 'dane_osobowe' in filepath_lower:
            return 'dane_osobowe'
        elif 'egzaminy' in filepath_lower:
            return 'egzaminy'
        elif 'rekrutacja' in filepath_lower:
            return 'rekrutacja'
        elif 'stypendia' in filepath_lower:
            return 'stypendia'
        elif 'urlopy_zwolnienia' in filepath_lower:
            return 'urlopy_zwolnienia'
        else:
            return 'general'
    
    #method: calculate MD5 hash of file
    def _calculate_file_hash(self, filepath: Path) -> str:
        """Calculate MD5 hash of file"""
        hash_md5 = hashlib.md5()
        with open(filepath, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()

# Singleton instance
document_processor = DocumentProcessor()