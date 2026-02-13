# agent4_bos/core/document_generator.py
import os
import time
from datetime import datetime
from docx import Document
from .config import KNOWLEDGE_BASE_PATH, KNOWLEDGE_BASE_CATEGORIES, ALL_CATEGORIES_KEY
from .llm_service import llm_service

class DocumentGenerator:
    def __init__(self):
        self.output_dir = KNOWLEDGE_BASE_PATH

    def generate_document(self, topic: str, category: str) -> dict:
        """
        Generates a DOCX document on the given topic in the specified category.
        Returns a dictionary with file information.
        """
        print(f"GENERATOR: Rozpoczynam generowanie dokumentu. Temat: '{topic}', Kategoria: '{category}'")
        
        # 1. Generowanie treści przez LLM
        content = self._generate_content_with_llm(topic)
        
        # 2. Utworzenie pliku DOCX
        file_info = self._create_docx_file(topic, content, category)
        
        return file_info

    def _generate_content_with_llm(self, topic: str) -> dict:
        """Generates the title, filename and body of the document using LLM based on the provided topic."""

        prompt = f"""Jesteś doświadczonym pracownikiem administracji uczelnianej.
Użytkownik prosi o przygotowanie dokumentu na podstawie opisu: "{topic}"

Twoim zadaniem jest:
1. Zrozumieć intencję użytkownika i rodzaj potrzebnego dokumentu.
2. Zaproponować profesjonalną nazwę pliku (krótką, bez polskich znaków, użyj podkreśleń zamiast spacji, np. Podanie_o_urlop).
3. Nadać dokumentowi oficjalny tytuł z prefiksem "AI_GEN_" (np. "AI_GEN_Podanie_o_urlop").
4. Przygotować kompletną treść dokumentu.

Twoja odpowiedź MUSI być w formacie:

NAZWA_PLIKU: [Nazwa_pliku_z_podkresleniami]
TYTUŁ: [Oficjalny Tytuł Dokumentu]
TREŚĆ:
[Treść dokumentu...]
"""
        response = llm_service.generate_response(prompt, temperature=0.4, max_tokens=2500)
        
        # Robust parser - analiza linia po linii
        lines = response.strip().split('\n')
        suggested_filename = None
        title = None
        body_lines = []
        is_body = False
        
        for line in lines:
            if is_body:
                body_lines.append(line)
                continue
            
            clean_line = line.strip()
            if not clean_line:
                continue
                
            upper_line = clean_line.upper()
            if upper_line.startswith("NAZWA_PLIKU:"):
                suggested_filename = line.split(":", 1)[1].strip()
            elif upper_line.startswith("TYTUŁ:") or upper_line.startswith("TYTUL:"):
                title = line.split(":", 1)[1].strip()
            elif upper_line.startswith("TREŚĆ:") or upper_line.startswith("TRESC:"):
                is_body = True
            else:
                if suggested_filename or title:
                    is_body = True
                    body_lines.append(line)
        
        body = "\n".join(body_lines).strip()
        
        # Fallback - jeśli parsowanie się nie uda, użyj całej odpowiedzi jako treści
        if not body and not suggested_filename:
            body = response
            
        if not title:
            title = f"Dokument: {topic}"

        return {"title": title, "body": body, "suggested_filename": suggested_filename}

    def _create_docx_file(self, topic: str, content: dict, category: str) -> dict:
        """Tworzy fizyczny plik .docx i zapisuje go w odpowiednim folderze"""
        try:
            doc = Document()
            
            # Nagłówek
            doc.add_heading(content['title'], 0)
            
            # Metadane w dokumencie
            doc.add_paragraph(f"Data wygenerowania: {datetime.now().strftime('%Y-%m-%d')}")
            doc.add_paragraph(f"Kategoria: {category}")
            doc.add_paragraph("-" * 50)
            
            # Treść właściwa
            for paragraph in content['body'].split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            # Ustalenie ścieżki zapisu
            target_category = category if category in KNOWLEDGE_BASE_CATEGORIES else "dane_osobowe"
            if category == ALL_CATEGORIES_KEY:
                target_category = "dane_osobowe"
                
            target_folder = os.path.join(self.output_dir, target_category)
            os.makedirs(target_folder, exist_ok=True)
            
            # Generowanie bezpiecznej nazwy pliku
            base_name = content.get("suggested_filename")
            

            if not base_name:
                if content['title'] and content['title'] != f"Dokument: {topic}":
                    base_name = content['title']
                else:
                    base_name = topic
            

            safe_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in base_name)[:60]
            safe_name = safe_name.replace(" ", "_")
            

            while "__" in safe_name:
                safe_name = safe_name.replace("__", "_")
            safe_name = safe_name.strip("_")
            
            timestamp = int(time.time())
            filename = f"AI_GEN_{safe_name}_{timestamp}.docx"
            filepath = os.path.join(target_folder, filename)
            

            doc.save(filepath)
            print(f"GENERATOR: Zapisano plik: {filepath}")
            

            relative_path = f"/data/knowledge_base/{target_category}/{filename}"
            
            return {
                "name": filename,
                "path": filepath,
                "download_url": relative_path,
                "success": True
            }
            
        except Exception as e:
            print(f"GENERATOR ERROR: {e}")
            return {"success": False, "error": str(e)}

# Singleton
document_generator = DocumentGenerator()