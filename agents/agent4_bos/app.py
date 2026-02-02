import os
from contextlib import asynccontextmanager
from fastapi import FastAPI, Form, HTTPException
from pathlib import Path
from docx import Document
import uuid
import requests
from qdrant_client.models import VectorParams, Distance
from langchain_community.chat_models import ChatOllama
from utils.embeddings import embed
from utils.qdrant_client import client, COLLECTION
from query import natural_agent_answer
from ingest import ingest_folder as auto_ingest_on_startup
from datetime import datetime

# Environment
OLLAMA_URL = os.getenv("OLLAMA_URL", "http://ollama:11434")
EMBED_MODEL = os.getenv("EMBED_MODEL", "nomic-embed-text")
COLLECTION_DOCS = os.getenv("COLLECTION", "collection_documents")
SPECIAL_COLLECTION = os.getenv("SPECIAL_COLLECTION", "special_cases")
DRAFTS_PATH = "/qdrant_data/drafts"

# Paths
DOCS_PATH = "/qdrant_data/documents"
SPECIAL_CASES_PATH = "/qdrant_data/special_cases"

# ===== LIFESPAN MANAGER =====
@asynccontextmanager
async def lifespan(app: FastAPI):
    """Handle startup and shutdown events."""
    # Startup
    print("🚀 Starting FastAPI application...")
    
    # Ensure collections exist
    ensure_collection(COLLECTION_DOCS)
    ensure_collection(SPECIAL_COLLECTION)
    
    print("🔧 Running hybrid ingest on startup...")
    docs_count = auto_ingest_on_startup(DOCS_PATH, COLLECTION_DOCS)
    special_count = auto_ingest_on_startup(SPECIAL_CASES_PATH, SPECIAL_COLLECTION)
    print(f"📊 Ingested {docs_count} docs, {special_count} special cases")
    
    yield  # App runs here
    
    # Shutdown (optional)
    print("👋 Shutting down...")

# Create FastAPI app with lifespan
app = FastAPI(lifespan=lifespan)

# LLM
llm = ChatOllama(model="llama3", base_url=OLLAMA_URL)

# Ensure collections exist - function
def ensure_collection(name: str):
    """Create collection if it doesn't exist."""
    try:
        collections = client.get_collections().collections
        collection_names = [c.name for c in collections]
        
        if name not in collection_names:
            print(f"📦 Creating collection: {name}")
            client.recreate_collection(
                collection_name=name,
                vectors_config={
                    "default": {
                        "size": 768,  # nomic-embed-text dimension
                        "distance": "Cosine"
                    }
                }
            )
            print(f"✅ Collection created: {name}")
        else:
            print(f"📊 Collection exists: {name}")
    except Exception as e:
        print(f"❌ Error ensuring collection {name}: {e}")
        raise

# Ensure drafts directory exists
def ensure_drafts_directory():
    """Create drafts directory if it doesn't exist."""
    drafts_path = Path(DRAFTS_PATH)
    if not drafts_path.exists():
        drafts_path.mkdir(parents=True, exist_ok=True)
        drafts_path.chmod(0o775)
        print(f"📁 Created drafts directory: {DRAFTS_PATH}")

#--------ENDPOINTS--------#

@app.post("/run")
async def run(payload: dict):
    """
    Simple endpoint to query LLM with text, for testing or generating summaries.
    """
    task = payload.get("input", "")
    if not task:
        raise HTTPException(status_code=400, detail="Field 'input' is required")
    
    response = llm.invoke(f"Generate draft response or summary for this: {task}")
    return {"draft": response.content, "collection": COLLECTION_DOCS}

@app.post("/special_case")
async def add_special_case(
    tytul: str = Form(...),
    opis: str = Form(...),
    autor: str = Form(...),
    uwagi: str = Form("")
):
    """
    Submit a special case via form with hybrid ingest:
    - Saves DOCX in special_cases folder
    - Uses hybrid system to ingest to Qdrant (no duplicates)
    """
    # Generate file path
    file_id = str(uuid.uuid4())
    filename = f"{file_id}_{tytul[:30].replace(' ', '_')}.docx"
    filepath = Path(SPECIAL_CASES_PATH) / filename
    
    # Ensure directory exists
    filepath.parent.mkdir(parents=True, exist_ok=True)

    # Save DOCX
    doc = Document()
    doc.add_heading(tytul, level=1)
    doc.add_paragraph(f"Opis: {opis}")
    doc.add_paragraph(f"Autor: {autor}")
    if uwagi:
        doc.add_paragraph(f"Uwagi dodatkowe: {uwagi}")
    doc.save(filepath)
    # Ensure file is writable
    filepath.chmod(0o666)
    
    print(f"💾 Saved DOCX: {filename}")
    
    # ===== HYBRID INGEST =====
    metadata = {
        "tytul": tytul,
        "opis": opis,
        "autor": autor,
        "uwagi": uwagi,
        "form_submission_id": file_id,
        "submitted_at": datetime.now().isoformat(),
        "ingest_method": "form_submission"
    }
    
    try:
        from ingest import ingest_form_submission
        ingest_result = ingest_form_submission(filepath, metadata)
        
        print(f"📊 Ingest result: {ingest_result}")
        
        return {
            "status": "ok",
            "file": filename,
            "id": file_id,
            "ingest_result": ingest_result,
            "message": "File saved and ingested to Qdrant" 
                if ingest_result.get("status") == "ingested" 
                else f"File saved ({ingest_result.get('reason', 'unknown')})"
        }
    
    except Exception as e:
        print(f"❌ Hybrid ingest failed: {e}")
        
        # Fallback: old method
        text = f"{tytul}\n{opis}\n{autor}\n{uwagi}"
        vector = embed(text)
        
        client.upsert(
            collection_name=SPECIAL_COLLECTION,
            points=[{
                "id": file_id,
                "vector": vector.tolist() if hasattr(vector, "tolist") else list(vector),
                "payload": {
                    "file_name": filename,
                    "tytul": tytul,
                    "opis": opis,
                    "autor": autor,
                    "uwagi": uwagi,
                    "ingest_method": "fallback"
                }
            }]
        )
        
        return {
            "status": "ok",
            "file": filename,
            "id": file_id,
            "ingest_result": {"status": "ingested", "method": "fallback"},
            "message": "File saved (using fallback ingest)"
        }

@app.post("/rag/query")
async def rag_query(payload: dict):
    """
    Main RAG endpoint using natural agent.
    """
    pytanie = payload.get("pytanie")
    if not pytanie:
        raise HTTPException(status_code=400, detail="Pole 'pytanie' jest wymagane")
    
    print(f"🧠 Processing query: '{pytanie}'")
    
    # Use natural agent
    result = natural_agent_answer(pytanie)
    
    # Optionally save draft
    save_draft = payload.get("save_draft", False)
    
    if save_draft and result.get("odpowiedz"):
        try:
            ensure_drafts_directory()
            filename = save_draft_docx(pytanie, result["odpowiedz"])
            result["plik_draft"] = filename
        except Exception as e:
            print(f"⚠️ Could not save draft: {e}")
            result["draft_error"] = str(e)
    
    return result

def save_draft_docx(title: str, content: str) -> str:
    """
    Save generated draft as DOCX in qdrant_data/drafts.
    """
    ensure_drafts_directory()
    
    file_id = str(uuid.uuid4())
    safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in title[:30])
    filename = f"DRAFT_{file_id}_{safe_title}.docx"
    filepath = Path(DRAFTS_PATH) / filename

    doc = Document()
    doc.add_heading(f"DRAFT – {title}", level=1)
    doc.add_paragraph(content)
    doc.save(filepath)
    
    # Set permissions
    filepath.chmod(0o664)
    
    print(f"💾 Saved draft: {filename}")
    return filename

# ===== ADMIN & DIAGNOSTIC ENDPOINTS =====

@app.get("/admin/status")
async def admin_status():
    """Get system status and ingest information."""
    try:
        docs_info = client.get_collection(COLLECTION_DOCS)
        special_info = client.get_collection(SPECIAL_COLLECTION)
        
        # Count files in folders
        docs_files = list(Path(DOCS_PATH).iterdir()) if Path(DOCS_PATH).exists() else []
        special_files = list(Path(SPECIAL_CASES_PATH).iterdir()) if Path(SPECIAL_CASES_PATH).exists() else []
        draft_files = list(Path(DRAFTS_PATH).iterdir()) if Path(DRAFTS_PATH).exists() else []
        
        return {
            "status": "healthy",
            "qdrant": {
                "documents_collection": {
                    "name": COLLECTION_DOCS,
                    "points": docs_info.points_count,
                    "vectors": docs_info.vectors_count
                },
                "special_cases_collection": {
                    "name": SPECIAL_COLLECTION,
                    "points": special_info.points_count,
                    "vectors": special_info.vectors_count
                }
            },
            "file_system": {
                "documents_folder": {
                    "path": DOCS_PATH,
                    "exists": Path(DOCS_PATH).exists(),
                    "files": len(docs_files),
                    "sample": [f.name for f in docs_files[:5]]
                },
                "special_cases_folder": {
                    "path": SPECIAL_CASES_PATH,
                    "exists": Path(SPECIAL_CASES_PATH).exists(),
                    "files": len(special_files),
                    "sample": [f.name for f in special_files[:5]]
                },
                "drafts_folder": {
                    "path": DRAFTS_PATH,
                    "exists": Path(DRAFTS_PATH).exists(),
                    "files": len(draft_files)
                }
            }
        }
    
    except Exception as e:
        return {
            "status": "error",
            "error": str(e)
        }

@app.post("/admin/ingest/sync")
async def sync_ingest():
    """Manually trigger sync of new files."""
    try:
        results = auto_ingest_on_startup()
        return {
            "status": "success",
            "message": "Ingest completed",
            "results": results
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ingest failed: {str(e)}")

@app.post("/admin/ingest/check-file")
async def check_file_in_qdrant(payload: dict):
    """Check if a specific file is in Qdrant."""
    from ingest import check_file_in_qdrant
    
    filename = payload.get("filename")
    collection = payload.get("collection", SPECIAL_COLLECTION)
    
    if not filename:
        raise HTTPException(status_code=400, detail="Field 'filename' is required")
    
    file_path = Path(SPECIAL_CASES_PATH) / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"File not found: {filename}")
    
    is_in_qdrant = check_file_in_qdrant(file_path, collection)
    
    return {
        "filename": filename,
        "collection": collection,
        "in_qdrant": is_in_qdrant,
        "file_exists": file_path.exists(),
        "file_size": file_path.stat().st_size if file_path.exists() else 0
    }

# Health check endpoint
@app.get("/health")
async def health_check():
    """Simple health check."""
    return {
        "status": "ok",
        "timestamp": datetime.now().isoformat(),
        "service": "agent4_bos",
        "version": "1.0"
    }

# Root endpoint
@app.get("/")
async def root():
    """Root endpoint with API information."""
    return {
        "service": "University Administrative Assistant",
        "endpoints": {
            "rag_query": "POST /rag/query - Ask questions",
            "special_case": "POST /special_case - Submit special cases",
            "run": "POST /run - Direct LLM queries",
            "admin_status": "GET /admin/status - System status",
            "health": "GET /health - Health check"
        },
        "documentation": "/docs or /redoc"
    }